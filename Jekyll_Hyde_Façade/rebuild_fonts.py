#!/usr/bin/env python3
"""Rebuild Façade.docx with Noto fonts. English song, Chinese annotations use SC."""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

FONT_SC = "Noto Serif CJK SC"
FONT_SANS = "Noto Sans CJK SC"
FONT_EN = "Arial"
DEEP_BLUE = "2C3E50"
AMBER = "B85C3A"
GOLD_BORDER = "D4A574"
LIGHT_BLUE = "D6E4F0"
WARM_BEIGE = "F5EBE0"
LIGHT_GREEN = "E8F0E4"
LIGHT_GRAY = "F8F8F8"

def set_run_font(run, cjk=FONT_SC, en=FONT_EN, size=Pt(11), bold=False, color=None):
    run.font.size = size
    run.bold = bold
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cjk)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_p(text, cjk=FONT_SC, en=FONT_EN, size=Pt(11), bold=False, color=None, alignment=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, cjk=cjk, en=en, size=size, bold=bold, color=color)
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    return p

def shade_cell(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{color}"/>')
    cell._element.get_or_add_tcPr().append(shading)

def make_table(headers, rows, header_color):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, header_color)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, cjk=FONT_SANS, en=FONT_EN, size=Pt(10), bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                shade_cell(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, cjk=FONT_SC, en=FONT_EN, size=Pt(9.5))
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()
    return table

def add_quote(text, size=Pt(10)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    pPr = p._element.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="12" w:color="{GOLD_BORDER}" w:space="8"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)
    run = p.add_run(text)
    set_run_font(run, cjk=FONT_SC, en=FONT_EN, size=size)
    run.italic = True
    return p

def add_grammar_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, cjk=FONT_SANS, en=FONT_EN, size=Pt(11), bold=True, color=AMBER)
    return p

# ============ BUILD ============

title_p = add_p("歌曲学习：Façade（假象）", cjk=FONT_SANS, en=FONT_EN, size=Pt(18), bold=True, color=DEEP_BLUE, space_after=Pt(4))
sep = doc.add_paragraph()
sep.paragraph_format.space_after = Pt(12)
sep_run = sep.add_run("─" * 40)
set_run_font(sep_run, cjk=FONT_SC, en=FONT_EN, size=Pt(8), color=DEEP_BLUE)

# == Basic Info ==
add_p("基本信息", cjk=FONT_SANS, en=FONT_EN, size=Pt(15), bold=True, space_after=Pt(8))

make_table(["项目", "内容"], [
    ["歌曲名", "Façade（法语，意为"外表/虚饰"）"],
    ["作曲", "Frank Wildhorn"],
    ["作词", "Leslie Bricusse（合写：Steve Cuden）"],
    ["所属作品", "音乐剧《Jekyll & Hyde》（变身怪医）"],
    ["首演年份", "1990年（Houston Alley Theatre）；1997年百老汇首演"],
    ["版本", "1997年百老汇原版录音（Original Broadway Cast Recording）"],
    ["语言", "英语"],
    ["类型", "群戏合唱（Ensemble Number），贯穿全剧的核心主题曲"],
    ["中文版译名", "《假象》（译配：周笑微）"],
], LIGHT_BLUE)

# == Background ==
add_p("背景故事", cjk=FONT_SANS, en=FONT_EN, size=Pt(15), bold=True, space_after=Pt(8))

add_p(""Façade" 是整部《变身怪医》音乐剧的主题锚点——它不仅是开场曲，还以 Reprise（重现段）的形式贯穿全剧，出现在订婚宴、红鼠妓院和大结局中，像希腊歌队一样持续评说维多利亚伦敦社会的伪善。")

add_p("Frank Wildhorn 早在南加州大学本科时代（1977年）就萌生了做这部音乐剧的念头。他说："大多数恐怖故事里，怪物来自外部。但在 Jekyll & Hyde 里，怪物来自内部——这对我们社会来说合适得多。"他和 Steve Cuden 在 80 年代末写出初稿，1988 年前后请来了 Les Misérables 级别的词作大师 Leslie Bricusse（曾执笔《Stop the World – I Want to Get Off》和《Goldfinger》）重写剧本和歌词。Bricusse 带来了一种尖锐的社会讽刺感，正是这种气质让"Façade"从普通的群唱变成了全剧的灵魂。")

add_p("Wildhorn 在 2000 年的一次采访中直接解释过这首歌："人们喜欢让世界看到的那一面，和关起门来真实的自己不是一回事。整部剧讲的就是瓶子里的妖怪——以及我们把它放出来之后会发生什么。"")

add_p("2013 年百老汇复排版的舞台设计将"Façade"的主题视觉化了：五位高层的董事会成员穿着内衣登场，仆人们一件一件地为他们穿戴礼服——社会身份本身就是一层一层搭建起来的"假象"。这个设计被剧评人称为整部剧最有力的舞台隐喻。")

add_p("从 1990 年首演至今，无论哪个版本，"Façade"从未被删减或替换——这在《变身怪医》这部歌曲阵容频繁变动的作品中几乎是唯一的。它不随版本更迭而动摇，因为它本身就是在说一件不会过时的事。")

# == Lyrics ==
add_p("歌词", cjk=FONT_SANS, en=FONT_EN, size=Pt(15), bold=True, space_after=Pt(4))
add_p("版本：1997年百老汇原版录音（Original Broadway Cast Recording），时长 3:54。歌词经 Genius、Musixmatch、Muzikum 三源交叉校验。", size=Pt(9), space_after=Pt(8))

lyrics_data = [
    ("There's a face that we wear / In the cold light of day — / It's society's mask, / It's society's way, / And the truth is / That it's all a façade!",
     "在白日冷光之下，我们戴着同一张脸——那是社会的面具，那是社会的规矩。而真相是：这一切都是假象！"),
    ("There's a face that we hide / Till the nighttime appears, / And what's hiding inside, / Behind all of our fears, / Is our true self, / Locked inside the façade!",
     "有一张脸我们藏到夜幕降临才敢露——藏在所有恐惧背后的，才是真正的自己，被锁在假象里面。"),
    ("Every day / People, in their own sweet way, / Like to add a coat of paint, / And be what they ain't! / That's how our little game is played, / Livin' like a masquerade, / Actin' a bizarre charade — / While playing the saint!",
     "每一天，人们用自己的方式，刷上新一层漆，装作自己不是的东西。这就是我们的小游戏：活成一场化装舞会，演一出荒诞的哑谜——同时假装圣人。"),
    ("But there's one thing I know, / And I know it for sure: / This disease that we've got / Has got no ready cure! / And I'm certain / Life is terribly hard — / When your life's a façade!",
     "但有件事我清楚，千真万确：我们得的这种病，没有现成的药方。而我可以肯定——当你的人生就是一场假象，活着真的很难。"),
    ("Look around you! I have found / You cannot tell, by lookin' at the surface, / What is lurkin' there beneath it! / See that face! Now, I'm prepared to bet you, / What you see's not what you get — / 'Cause man's a master of deceit!",
     "看看你周围！我发现——你没法光看表面就知道底下藏着什么。看那张脸！我敢跟你打赌，你看到的不是你想象的样子——因为人类是欺骗的大师！"),
    ("So, what is the sinister secret? / The lie he will tell you is true? — / It's that each man you meet / In the street / Isn't one man but two!",
     "那么，那个阴暗的秘密是什么？他对你说的谎言难道还是真的？——那就是你在街上遇到的每一个人，都不是一个人，而是两个！"),
    ("Nearly everyone you see — / Like him an' her, an' you, an' me — / Pretends to be / A pillar of society — / A model for propriety — / Sobriety an' piety — / Who shudders at the thought / Of notoriety!",
     "你看到的几乎每个人——就像他和她，和你，和我——都假装是社会栋梁、道德典范、清醒虔诚的好人，想到"声名狼藉"就假装发抖！"),
    ("The ladies an' gents 'ere before you — / Which none of 'em ever admits — / May 'ave saintly looks — / But they're sinners an' crooks! / Hypocrites! Hypocrites!",
     "站在你面前的各位先生女士——虽然没一个会承认——长得像圣人，实则是罪人和恶棍！伪君子！伪君子！"),
    ("There are preachers who kill! / There are killers who preach! / There are teachers who lie! / There are liars who teach! / Take yer pick, dear — / 'Cause it's all a façade!",
     "有会杀人的牧师！有会布道的杀手！有会说谎的教师！有会教书的骗子！亲爱的，随便选——因为这一切都是假象！"),
    ("If we're not one, but two, / Are we evil or good? / Do we walk the fine line — / That we'd cross if we could? / Are we waiting — / To break through the façade?",
     "如果我们不止一个，而是两个——我们是善还是恶？我们走在那条微妙的临界线上——如果能跨过去，我们就跨过去吗？我们是不是在等待——冲破这假象的机会？"),
    ("One or two / Might look kinda well-to-do — / Hah! They're bad as me an' you, / Right down they're boots! / I'm inclined to think — / Half mankind / Thinks the other half is blind! / Wouldn't be surprised to find — / They're all in cahoots!",
     "有一两个人看起来人模人样——哈！他们和你我一样坏，彻头彻尾！我倾向于认为——半个人类觉得另一半人类是瞎的！如果他们全都是一伙的，我也不意外！"),
    ("At the end of the day, / They don't mean what they say, / They don't say what they mean, / They don't ever come clean — / And the answer — / Is it's all a façade!",
     "说到底——他们说出来的话不真心，真心的话说不出来，从来不肯坦白——而答案就是：这一切都是假象！"),
    ("Man is not one, but two, / He is evil and good, / An' he walks the fine line / We'd all cross if we could! / It's a nightmare — / We can never discard — / So we stay on our guard — / Though we love the façade — / What's behind the façade? / Look behind the façade!",
     "人不是一个，而是两个——他既恶又善。他走在我们都想跨过的临界线上。这是一场永远摆脱不了的噩梦——所以我们永远警惕，尽管我们爱这假象。假象背后是什么？看看假象背后！"),
]

make_table(["原文", "中文翻译"], lyrics_data, WARM_BEIGE)

add_p("歌词中有意使用了法文词 façade（带变音符 ç）而非英文拼写 "facade"——这是 Bricusse 对手戏的精心设计：用法文词的"上流感"来包装对维多利亚上流社会的嘲讽。歌词中还有 masquerade、charade、sobriety、piety、propriety 等一连串押韵的"体面词汇"，堆叠出轰隆的讽刺效果。", size=Pt(9), space_after=Pt(12))

# == Language Learning ==
add_p("语言学习", cjk=FONT_SANS, en=FONT_EN, size=Pt(15), bold=True, space_after=Pt(8))
add_p("核心词汇", cjk=FONT_SANS, en=FONT_EN, size=Pt(13), bold=True, space_after=Pt(4))
add_p("从歌词中选出 12 个有学习价值的英文词/词组——有 GRE 级高阶词，也有音乐剧特有的口语化表达。", size=Pt(9), space_after=Pt(6))

vocab_data = [
    ["façade", "fə-SAHd（法语借词, ç 发 s 音）", "n.", "外表、假象；（建筑）立面", "★★★"],
    ["masquerade", "mas-kə-RAYD", "n./v.", "化装舞会；伪装", "★★☆"],
    ["charade", "shə-RAYD", "n.", "哑谜猜字游戏；明显的做戏/伪装", "★★☆"],
    ["propriety", "prə-PRY-ə-tee", "n.", "得体、合乎礼仪的行为", "★★★"],
    ["sobriety", "sə-BRY-ə-tee", "n.", "清醒、未醉；严肃持重", "★★★"],
    ["notoriety", "no-tə-RY-ə-tee", "n.", "声名狼藉、恶名", "★★☆"],
    ["be in cahoots (with)", "kə-HOOTS", "phr.", "勾结、共谋（口语）", "★★☆"],
    ["pillar of society", "—", "phr.", "社会栋梁（常带讽刺语气）", "★☆☆"],
    ["come clean", "—", "phr. v.", "坦白、说实话", "★☆☆"],
    ["ain't / an' / 'ere / yer / kinda", "—", "缩略", "歌词中大量使用口语缩略，模仿伦敦下层口音", "★☆☆"],
    ["fine line", "—", "n.", "细微的分界线（在两种状态间艰难平衡）", "★☆☆"],
    ["lurking / lurkin'", "LURK-ing", "v.", "潜伏、潜藏", "★★☆"],
]
make_table(["单词/短语", "音标/发音提示", "词性", "释义", "难度"], vocab_data, LIGHT_GREEN)

add_p("选词说明：Façade 作为歌名是全曲核心，理解它的发音和法源是关键。propriety/sobriety/notoriety 三词在歌词中并列押韵（-iety），形成 Bricusse 标志性的"轰炸式讽刺"。come clean 和 be in cahoots 是口语高频短语，音乐剧歌词常用口语体表达。", size=Pt(9), space_after=Pt(12))

# Grammar points
add_p("语法与修辞点", cjk=FONT_SANS, en=FONT_EN, size=Pt(13), bold=True, space_after=Pt(4))
add_p("从歌词中提取 4 个值得注意的语言现象。", size=Pt(9), space_after=Pt(8))

gp1_title = "法语借词 façade 的英文用法"
gp1_body = "façade 在英文中的使用频率远高于中文对"假象"的使用，它既可以指建筑学意义上的"（建筑物）正面/立面"，也可以比喻"伪装出来的外表"。歌词的巧妙在于，它在"face"和"façade"之间建立了音韵联系——先唱 "There's a face that we wear"，然后用 façade 作为概念的升级版（face → façade = 脸 → 精心构建的假面）。写作中可以用 "beneath the façade" 或 "behind the façade" 来表达"表象之下"。"

gp2_title = ""What you see's not what you get" —— 否定式习语的活用"
gp2_body = "标准谚语是 "What you see is what you get"（所见即所得，缩写 WYSIWYG 也成了计算机术语）。Bricusse 把它反转——所见并非所得——来表达人的不可信。这种反转既有文化共鸣（听众都认得原句），又精准服务于剧情：Hyde 的恐怖就在于他藏在"正常人"的外表下。"

gp3_title = ""They don't mean what they say / They don't say what they mean" —— Chiasmus（交错配列）"
gp3_body = "这是英文修辞中经典的 chiasmus 结构——"mean→say / say→mean"，AB-BA 对称。两句看似在说同一件事，实则互为镜像："不说真心话"和"说不出真话"是从不同方向剖开同一个问题。日常英文中你也可以用：I don't know what I want, and I don't want what I know."

gp4_title = "歌词中的小舌音/口语缩略"
gp4_body = "全曲大量使用 ain't / an' / 'ere / 'em / lookin' / lurkin' / kinda / yer 等拼写缩略，这不是打字错误——这是在用拼写暗示角色口音。音乐剧中这叫"eye dialect"（视觉方言），用非标准拼写告诉演员"用下层口音唱"。对于英语学习者，重点不是模仿这种口音，而是能识别它，知道原形。ain't 在正式写作中避免使用，但在歌曲、对话中理解它的语法功能就够了。"

for t, b in [(gp1_title, gp1_body), (gp2_title, gp2_body), (gp3_title, gp3_body), (gp4_title, gp4_body)]:
    add_grammar_title(t)
    add_p(b, size=Pt(10.5), space_after=Pt(8))

# Culture notes
add_p("文化笔记", cjk=FONT_SANS, en=FONT_EN, size=Pt(13), bold=True, space_after=Pt(6))

cn1_title = "维多利亚伦敦的"假象"——不只是剧情设定，是历史现实"
cn1_body = "1886 年 Stevenson 写《化身博士》时，伦敦正处于工业革命巅峰期——表面上是大英帝国最辉煌的年代，实际上是贫富差距极端撕裂、犯罪率飙升、东区贫民窟连警察都不敢进的时代。所以"Façade"唱的社会虚伪不是虚构——中产阶级上教堂、穷人捡垃圾、同一批"社会栋梁"白天在董事会装君子、晚上去红鼠妓院。中文版译名《假象》也抓住了这点：不是"面具"那种主动伪装，而是"表象"本身就在欺骗。"

cn2_title = "为什么用 "façade" 这个法文词？"
cn2_body = "维多利亚时代的英国上流社会以说法语为教养标志（法语是欧洲贵族和外交的通用语）。Bricusse 选择法文拼写的"façade"而非英文"facade"是用上流社会的语言来描述上流社会的虚伪——这个选择本身就是嘲讽。ç 这个字母在英文里不存在，但在法文中是软音符号（cedilla），表示 s 的音——而英文原本是可以直接拼"fasade"的。对英语母语者来说，看到 ç 就会有"外来、高级、不自然"的直觉——这就是歌词要传达的。"

cn3_title = ""Isn't one man but two" 的神学背景"
cn3_body = "这句是全剧最核心的台词。Stevenson 的小说发表于 1886 年，同一时期达尔文进化论（1859年）正在颠覆基督教世界的"人按上帝形象造"的信念，弗洛伊德也在酝酿潜意识理论。Jekyll & Hyde 的故事出现在这个断裂带上——人不只是一个统一的好灵魂，而是一个分裂的、有黑暗驱力的动物。"Man is not one, but two"既是对剧情的预告，也是对维多利亚人自我认知的挑战。"

for t, b in [(cn1_title, cn1_body), (cn2_title, cn2_body), (cn3_title, cn3_body)]:
    add_grammar_title(t)
    add_p(b, size=Pt(10.5), space_after=Pt(8))

# == Singing techniques ==
add_p("演唱技巧", cjk=FONT_SANS, en=FONT_EN, size=Pt(15), bold=True, space_after=Pt(8))

tips = [
    ("全团 SATB 和声的紧凑度",
     ""Façade" 是 SATB（女高/女中/男高/男低）编制的群唱，全剧最高音达到 High C (C6)，出现在女高音声部。",
     "最高音不在炫技段，而在"Façade!"的句末强音上——声音不能飘，要靠横膈膜强力支撑。合唱团四个声部之间不能有一丁点错位，因为乐句密集、节奏脉冲式推进。"),
    ("Belt 与混合声的快速切换",
     "Frank Wildhorn 的风格混合了摇滚和百老汇，这首歌尤其如此。段落从低声叙述（"There's a face that we wear"）突跳到全团爆发（"It's all a façade!"），要求胸腔强声（belt）的自然爆发。",
     "不要在中低声区就提前蓄力——把能量留给"façade"这个锚点词。爆发时下颌放松、保持咽腔打开，不要压缩喉咙。"),
    ("密集歌词的咬字精度",
     ""Every day / People in their own sweet way / Like to add a coat of paint / And be what they ain't"——这段每一拍都有词，而且大量辅音（p/t/k 塞音集中）。",
     "用节拍器慢速练习舌头位置，每个"t"用舌尖点硬腭、不送气（这点和日语的 t 有类似之处，但英文的送气略强）。先在说唱状态过一遍，再回到旋律。"),
    ("不和谐和弦的氛围营造",
     "开头以不和谐和弦切入，合唱团从乐池中传出声响，营造一种诡异、不安的感觉。",
     "指挥通常会在开头的"cold light of day"处压低声部，不让声音太"漂亮"——需要一种苍白、冷酷的质感。有的制作版本在 Reprise 段用全耳语演唱，创造如履薄冰的紧张感。"),
    (""Hypocrites!"——全曲最难的爆发点",
     "两个连续的"Hypocrites!"，四声部齐吼，节奏切分激烈。",
     "这个段落不是"唱"出来的，是"砸"出来的。气息需要在前面一句"but they're sinners and crooks"的尾音就预吸好，"Hy-"开头的/h/音不要呼出太多气——用短促的一击。"),
    ("最后一段的三重递进",
     "结尾 "What's behind the façade? / Look behind the façade!" 是三个层次的收束——从询问到命令，从怀疑到揭示。",
     ""What's behind"用试探的语气，"Look behind"直接转向观众。最后一个词的"çade"要留出空间：不是渐弱到无，而是在高潮位骤然收住，让观众在静默中消化全曲的质问。"),
]

for tip_title, tip_problem, tip_solution in tips:
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(8)
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run(tip_title)
    set_run_font(r_title, cjk=FONT_SANS, en=FONT_EN, size=Pt(11), bold=True)

    p_problem = doc.add_paragraph()
    p_problem.paragraph_format.left_indent = Cm(0.3)
    r_problem = p_problem.add_run(tip_problem)
    set_run_font(r_problem, cjk=FONT_SC, en=FONT_EN, size=Pt(10.5))

    p_solution = doc.add_paragraph()
    p_solution.paragraph_format.left_indent = Cm(0.6)
    p_solution.paragraph_format.space_after = Pt(6)
    r_solution = p_solution.add_run("→ " + tip_solution)
    set_run_font(r_solution, cjk=FONT_SC, en=FONT_EN, size=Pt(10.5))

# Separator
doc.add_paragraph()
sep2 = doc.add_paragraph()
sep2_run = sep2.add_run("─" * 40)
set_run_font(sep2_run, cjk=FONT_SC, en=FONT_EN, size=Pt(8), color=DEEP_BLUE)

# == Appendix ==
add_p("附录", cjk=FONT_SANS, en=FONT_EN, size=Pt(15), bold=True, space_after=Pt(8))

add_p("歌词来源", cjk=FONT_SANS, en=FONT_EN, size=Pt(13), bold=True, space_after=Pt(4))
for s in [
    "Genius — Frank Wildhorn ft. Anthony Warlow & Concept Cast: https://genius.com/Frank-wildhorn-facade-lyrics",
    "Musixmatch — Jekyll & Hyde Original Broadway Cast: https://www.musixmatch.com/lyrics/Jekyll-Hyde-The-Musical-Original-Broadway-Cast/Façade",
    "Muzikum — Frank Wildhorn Facade lyrics: https://muzikum.eu/en/frank-wildhorn/jekyll-and-hyde-the-original-broadway-production-facade-lyrics",
    "MTI (Music Theatre International) — Jekyll & Hyde Show History: https://www.mtishows.com/show-history/1049",
]:
    add_p(s, size=Pt(9), space_after=Pt(2))

add_p("背景信息来源", cjk=FONT_SANS, en=FONT_EN, size=Pt(13), bold=True, space_after=Pt(4))
for s in [
    "Frank Wildhorn interview — Salt Lake Tribune (2000): https://newspapers.lib.utah.edu/ark:/87278/s6g49z41/27797036",
    ""Jekyll composer finally realizes his double vision" — The Virginian-Pilot (1996)",
    "Jekyll & Hyde Wikipedia: https://en.wikipedia.org/wiki/Jekyll_%26_Hyde_(musical)",
    "SingingCarrots — Jekyll & Hyde vocal range: https://singingcarrots.com/artist-range?artist=Jekyll+%26+Hyde",
]:
    add_p(s, size=Pt(9), space_after=Pt(2))

add_p("参考音视频", cjk=FONT_SANS, en=FONT_EN, size=Pt(13), bold=True, space_after=Pt(4))
for s in [
    "1997 Broadway Cast Recording, Track 3 "Façade" — 时长 3:54，指挥 Jason Howland",
    "2013 百老汇复排版 — 开场群戏视觉化"穿衣"概念",
    "台大合唱团 纯人声翻唱: https://www.bilibili.com/video/BV1DN4y1L7Yu/",
    "阿卡主义人声乐团 中文版《假象》: https://www.bilibili.com/video/BV16x411n7Hc/",
]:
    add_p(s, size=Pt(9), space_after=Pt(2))

# Verify block
doc.add_paragraph()
verify_p = doc.add_paragraph()
verify_p.paragraph_format.space_before = Pt(12)
pPr = verify_p._element.get_or_add_pPr()
borders = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:left w:val="single" w:sz="12" w:color="{DEEP_BLUE}" w:space="8"/>'
    f'</w:pBdr>'
)
pPr.append(borders)
shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="F8F8F8"/>')
pPr.append(shading)

v_text = "歌词校验状态：英文原词经 Genius / Musixmatch / Muzikum 三源交叉验证一致。中文翻译独立译出，参照了周笑微译配版《假象》的部分措辞但保持直译优先原则。Façade 的变音符 ç 在部分歌词网站丢失（显示为 Facade），本文统一保留法文原拼写。"
v_run = verify_p.add_run(v_text)
set_run_font(v_run, cjk=FONT_SC, en=FONT_EN, size=Pt(9))

# Save
out_path = "E:/song-study/Jekyll_Hyde_Façade/Jekyll_Hyde_Façade.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
print(f"Size: {os.path.getsize(out_path)} bytes")
