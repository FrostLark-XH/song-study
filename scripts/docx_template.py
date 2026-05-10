#!/usr/bin/env python3
"""
可复用的 Word 文档生成器 —— song-study skill 专用。
生成带色彩语义系统的 .docx 文件（基本信息/歌词/词汇/语法/演唱技巧/校验声明）。

用法：作为参考模板，根据具体歌曲数据填充 Content 部分后运行。
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ═══ 色彩常量 ═══
DEEP_BLUE = "2C3E50"
LIGHT_BLUE = "D6E4F0"
WARM_BEIGE = "F5EBE0"
LIGHT_GREEN = "E8F0E4"
WARM_GOLD = "FFF3E0"
AMBER = "B85C3A"
LIGHT_GRAY = "F8F8F8"
BORDER_GOLD = "D4A574"


def set_font(run, cn="SimSun", en="Arial", size=Pt(11), bold=False, color=None):
    """设置 run 的字体（中/英/日文分别指定）。"""
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {
        qn('w:eastAsia'): cn,
        qn('w:ascii'): en,
        qn('w:hAnsi'): en,
    })
    rPr.insert(0, rFonts)


def set_cell_shading(cell, color):
    tcPr = cell._element.get_or_add_tcPr()
    shading = tcPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color,
    })
    tcPr.append(shading)


def make_table(doc, headers, rows, header_color, col_widths=None):
    """创建带表头着色和交替行着色的表格。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(h)
        set_font(run, cn="Microsoft YaHei", size=Pt(10), bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(val)
            set_font(run, size=Pt(10))
            if i % 2 == 1:
                set_cell_shading(cell, LIGHT_GRAY)
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Cm(w)
    doc.add_paragraph()
    return table


def add_quote_paragraph(doc, text):
    """歌手原话/创作引文：左缩进 + 暖金左边框 + 斜体。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    left = pBdr.makeelement(qn('w:left'), {
        qn('w:val'): 'single', qn('w:sz'): '6',
        qn('w:space'): '5', qn('w:color'): BORDER_GOLD,
    })
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    set_font(run, size=Pt(11))
    run.italic = True


def add_verification_block(doc, text):
    """文档末尾校验声明块：浅灰底色 + 深蓝灰左边框。"""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:fill'): LIGHT_GRAY,
    })
    pPr.append(shd)
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    left = pBdr.makeelement(qn('w:left'), {
        qn('w:val'): 'single', qn('w:sz'): '12',
        qn('w:space'): '8', qn('w:color'): DEEP_BLUE,
    })
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    set_font(run, size=Pt(10), bold=True, color=DEEP_BLUE)


def add_singing_tip(doc, tip_name, problem, solution):
    """演唱技巧条目：暖金底色包裹。"""
    for text, indent, bold in [
        (tip_name, Cm(0), True),
        (problem, Cm(0.3), False),
        ("→ " + solution, Cm(0.8), False),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = indent
        p.paragraph_format.space_after = Pt(2) if indent > Cm(0) else Pt(6)
        pPr = p._element.get_or_add_pPr()
        shd = pPr.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear', qn('w:fill'): WARM_GOLD,
        })
        pPr.append(shd)
        run = p.add_run(text)
        set_font(run, size=Pt(11) if bold else Pt(10.5), bold=bold)


def create_document(output_path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    return doc


# ═══ 以下是文档内容构建函数，按需调用 ═══
# 使用 create_document(path) 创建 doc 对象
# 使用 make_table(doc, ...) 创建表格
# 使用 add_quote_paragraph(doc, text) 添加引文
# 使用 add_singing_tip(doc, name, problem, solution) 添加演唱技巧
# 使用 add_verification_block(doc, text) 添加校验声明
# 最后 doc.save(output_path)
