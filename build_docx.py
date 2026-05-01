#!/usr/bin/env python3
"""Read song-study markdown files and generate .docx with system fonts.
Chinese: SimSun (body) / Microsoft YaHei (heading)
Japanese: MS Mincho (body) / Meiryo (heading)
English: Arial
"""
import sys, os, re, json
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# FONT CONFIG - system fonts guaranteed on Windows
# ============================================================
FONT_SC_BODY = "SimSun"
FONT_SC_HEAD = "Microsoft YaHei"
FONT_JP_BODY = "MS Mincho"
FONT_JP_HEAD = "Meiryo"
FONT_EN = "Arial"

# Colors
DEEP_BLUE = "2C3E50"
AMBER = "B85C3A"
GOLD_BORDER = "D4A574"
LIGHT_BLUE = "D6E4F0"
WARM_BEIGE = "F5EBE0"
LIGHT_GREEN = "E8F0E4"
LIGHT_GRAY = "F8F8F8"

# ============================================================
# DOCX HELPERS
# ============================================================
def set_font(run, cjk, en=FONT_EN, size=Pt(11), bold=False, color=None):
    run.font.size = size
    run.bold = bold
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml('<w:rFonts ' + nsdecls("w") + ' />')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cjk)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_p(doc, text, cjk=FONT_SC_BODY, en=FONT_EN, size=Pt(11), bold=False, color=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, cjk=cjk, en=en, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = space_after
    return p

def shade_cell(cell, color):
    shading = parse_xml('<w:shd ' + nsdecls("w") + ' w:val="clear" w:fill="' + color + '"/>')
    cell._element.get_or_add_tcPr().append(shading)

def make_table(doc, headers, rows, header_color, data_cjk=FONT_SC_BODY, header_cjk=FONT_SC_HEAD):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, header_color)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_font(run, cjk=header_cjk, en=FONT_EN, size=Pt(10), bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                shade_cell(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_font(run, cjk=data_cjk, en=FONT_EN, size=Pt(9.5))
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()
    return table

def add_quote(doc, text, cjk=FONT_SC_BODY, size=Pt(10)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    pPr = p._element.get_or_add_pPr()
    borders = parse_xml(
        '<w:pBdr ' + nsdecls("w") + '>'
        '<w:left w:val="single" w:sz="12" w:color="' + GOLD_BORDER + '" w:space="8"/>'
        '</w:pBdr>'
    )
    pPr.append(borders)
    run = p.add_run(text)
    set_font(run, cjk=cjk, en=FONT_EN, size=size)
    run.italic = True
    return p

def add_grammar_title(doc, text, cjk=FONT_SC_HEAD):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, cjk=cjk, en=FONT_EN, size=Pt(11), bold=True, color=AMBER)
    return p

def add_singing_tip(doc, title, problem, solution, title_cjk=FONT_SC_HEAD, body_cjk=FONT_SC_BODY):
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(8)
    p_title.paragraph_format.space_after = Pt(2)
    set_font(p_title.add_run(title), cjk=title_cjk, en=FONT_EN, size=Pt(11), bold=True)

    p_prob = doc.add_paragraph()
    p_prob.paragraph_format.left_indent = Cm(0.3)
    set_font(p_prob.add_run(problem), cjk=body_cjk, en=FONT_EN, size=Pt(10.5))

    p_sol = doc.add_paragraph()
    p_sol.paragraph_format.left_indent = Cm(0.6)
    p_sol.paragraph_format.space_after = Pt(6)
    set_font(p_sol.add_run("→ " + solution), cjk=body_cjk, en=FONT_EN, size=Pt(10.5))

def add_separator(doc):
    doc.add_paragraph()
    sep = doc.add_paragraph()
    sep.paragraph_format.space_after = Pt(12)
    r = sep.add_run("─" * 40)
    set_font(r, cjk=FONT_SC_BODY, en=FONT_EN, size=Pt(8), color=DEEP_BLUE)

def add_verify_block(doc, text, cjk=FONT_SC_BODY):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    pPr = p._element.get_or_add_pPr()
    borders = parse_xml(
        '<w:pBdr ' + nsdecls("w") + '>'
        '<w:left w:val="single" w:sz="12" w:color="' + DEEP_BLUE + '" w:space="8"/>'
        '</w:pBdr>'
    )
    pPr.append(borders)
    shading = parse_xml('<w:shd ' + nsdecls("w") + ' w:val="clear" w:fill="F8F8F8"/>')
    pPr.append(shading)
    set_font(p.add_run(text), cjk=cjk, en=FONT_EN, size=Pt(9))

# ============================================================
# BUILD DOCX FROM JSON DATA
# ============================================================
def build_docx(json_path, out_path, is_japanese_song=True):
    with open(json_path, 'r', encoding='utf-8') as f:
        d = json.load(f)

    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Font selection based on language
    sc_body = d.get('font_cn_body', FONT_SC_BODY)
    sc_head = d.get('font_cn_heading', FONT_SC_HEAD)
    jp_body = d.get('font_jp_body', FONT_JP_BODY)
    jp_head = d.get('font_jp_heading', FONT_JP_HEAD)

    # Title
    add_p(doc, d['title'], cjk=sc_head, size=Pt(18), bold=True, color=DEEP_BLUE, space_after=Pt(4))
    add_separator(doc)

    # Basic Info
    add_p(doc, '基本信息', cjk=sc_head, size=Pt(15), bold=True, space_after=Pt(8))
    make_table(doc, ['项目', '内容'], d['info_rows'], LIGHT_BLUE, data_cjk=sc_body, header_cjk=sc_head)

    # Background
    add_p(doc, '背景故事', cjk=sc_head, size=Pt(15), bold=True, space_after=Pt(8))
    for i, para in enumerate(d['bg_paras']):
        add_p(doc, para, cjk=sc_body, size=Pt(11), space_after=Pt(6))

    # Quote blocks (Japanese song special)
    if 'quote_jp' in d:
        add_p(doc, '', space_after=Pt(2))
        add_quote(doc, d['quote_jp'], cjk=jp_body)
        add_p(doc, '', space_after=Pt(2))
        add_quote(doc, d['quote_cn'], cjk=sc_body)
        add_p(doc, '', space_after=Pt(4))

    # Detect song type from lyric structure
    lyric_cols = 3
    if d['lyric_sections']:
        first_lines = d['lyric_sections'][0][1]
        if first_lines and len(first_lines[0]) == 2:
            lyric_cols = 2

    # Detect if vocab is Japanese or English
    is_jp_vocab = False
    if d['vocab_rows'] and d['vocab_rows'][0]:
        first_word = d['vocab_rows'][0][0]
        # Check if contains hiragana/katakana/kanji outside BMP
        is_jp_vocab = any('぀' <= c <= 'ヿ' or '一' <= c <= '鿿' for c in first_word)

    # Lyrics
    add_p(doc, '歌词', cjk=sc_head, size=Pt(15), bold=True, space_after=Pt(4))
    if is_jp_vocab:
        add_p(doc, '✅ 歌词来源：日文原文 + 罗马音 -> Genius Japan。中文翻译 -> 基于日文原文 + Genius 英文翻译综合译出。', cjk=sc_body, size=Pt(9), space_after=Pt(8))
    else:
        add_p(doc, '✅ 歌词来源：经 Genius、Musixmatch、Muzikum 三源交叉校验。', cjk=sc_body, size=Pt(9), space_after=Pt(8))

    for section_name, lines in d['lyric_sections']:
        if section_name:
            add_p(doc, section_name, cjk=sc_head, size=Pt(11), bold=True, color=DEEP_BLUE, space_after=Pt(4))
        if lyric_cols == 3:
            make_table(doc, ['原文', '发音（罗马字）', '中文翻译'], lines, WARM_BEIGE, data_cjk=jp_body, header_cjk=sc_head)
        else:
            make_table(doc, ['原文', '中文翻译'], lines, WARM_BEIGE, data_cjk=sc_body, header_cjk=sc_head)

    # Language Learning
    add_p(doc, '语言学习', cjk=sc_head, size=Pt(15), bold=True, space_after=Pt(8))

    # Vocab
    add_p(doc, '核心词汇', cjk=sc_head, size=Pt(13), bold=True, space_after=Pt(4))
    add_p(doc, '从全曲选出 ' + str(len(d['vocab_rows'])) + ' 个有学习价值的词，优先日常高频词和歌词核心意象词。', cjk=sc_body, size=Pt(9), space_after=Pt(6))
    if is_jp_vocab:
        make_table(doc, ['原文', '读音', '词性', '释义', '等级'], d['vocab_rows'], LIGHT_GREEN, data_cjk=jp_body, header_cjk=sc_head)
    else:
        make_table(doc, ['单词/短语', '音标/发音提示', '词性', '释义', '难度'], d['vocab_rows'], LIGHT_GREEN, data_cjk=sc_body, header_cjk=sc_head)

    if 'vocab_note' in d and d['vocab_note']:
        add_p(doc, d['vocab_note'], cjk=sc_body, size=Pt(9), space_after=Pt(4))
    if 'verb_note' in d and d['verb_note']:
        add_p(doc, d['verb_note'], cjk=sc_body, size=Pt(9), space_after=Pt(12))

    # Grammar
    add_p(doc, '语法点', cjk=sc_head, size=Pt(13), bold=True, space_after=Pt(4))
    add_p(doc, '从歌词中提取 ' + str(len(d['grammar_points'])) + ' 个有价值的句型/表达。', cjk=sc_body, size=Pt(9), space_after=Pt(8))
    for gp_title, gp_body in d['grammar_points']:
        add_grammar_title(doc, gp_title, cjk=sc_head)
        add_p(doc, gp_body, cjk=sc_body, size=Pt(10.5), space_after=Pt(8))

    # Culture
    add_p(doc, '文化笔记', cjk=sc_head, size=Pt(13), bold=True, space_after=Pt(6))
    for cn_title, cn_body in d['culture_notes']:
        add_grammar_title(doc, cn_title, cjk=sc_head)
        add_p(doc, cn_body, cjk=sc_body, size=Pt(10.5), space_after=Pt(8))

    # Singing Tips
    add_p(doc, '演唱技巧', cjk=sc_head, size=Pt(15), bold=True, space_after=Pt(8))
    for tip_title, tip_problem, tip_solution in d['singing_tips']:
        add_singing_tip(doc, tip_title, tip_problem, tip_solution, title_cjk=sc_head, body_cjk=sc_body)

    # Appendix
    add_separator(doc)
    add_p(doc, '附录', cjk=sc_head, size=Pt(15), bold=True, space_after=Pt(8))
    add_p(doc, '歌词来源', cjk=sc_head, size=Pt(13), bold=True, space_after=Pt(4))
    for s in d['sources_lyrics']:
        add_p(doc, s, cjk=sc_body, size=Pt(9), space_after=Pt(2))
    add_p(doc, '背景信息来源', cjk=sc_head, size=Pt(13), bold=True, space_after=Pt(4))
    for s in d['sources_bg']:
        add_p(doc, s, cjk=sc_body, size=Pt(9), space_after=Pt(2))

    # Verify
    add_verify_block(doc, d['verify_text'], cjk=sc_body)

    # Save
    doc.save(out_path)
    print('OK: ' + out_path + ' (' + str(os.path.getsize(out_path)) + ' bytes)')


# ============ MAIN ============
if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: python build_docx.py <json_data_file>')
        print('  or: python build_docx.py --all')
        sys.exit(1)

    if sys.argv[1] == '--all':
        # Build all known songs
        songs = [
            (r'E:\song-study\tuki._騙シ愛\data.json', r'E:\song-study\tuki._騙シ愛\tuki._騙シ愛.docx'),
            (r'E:\song-study\Jekyll_Hyde_Façade\data.json', r'E:\song-study\Jekyll_Hyde_Façade\Jekyll_Hyde_Façade.docx'),
        ]
        for json_path, out_path in songs:
            if os.path.exists(json_path):
                build_docx(json_path, out_path)
            else:
                print('MISSING: ' + json_path)
    else:
        build_docx(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else sys.argv[1].replace('.json', '.docx'))
